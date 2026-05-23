import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import subprocess

TEMPLATES_DIR = "generated_docs"
os.makedirs(TEMPLATES_DIR, exist_ok=True)


def generate_ijara_shartnoma(data: dict, order_id: int) -> str:
    """Ijara shartnomasi generatsiya qilish"""
    doc = Document()

    # Style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Sarlavha
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KO'CHMAS MULKNI VAQTINCHA FOYDALANISHGA BERISH\nSHARTNOMASI")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Sana va joy
    p = doc.add_paragraph()
    p.add_run(f"{data.get('shahar', 'Toshkent')} shahri").bold = True
    p.add_run(f"\t\t\t\t{data.get('sana', datetime.now().strftime('%d.%m.%Y'))} yil")

    doc.add_paragraph()

    # Tomonlar
    doc.add_paragraph(
        f"Bir tomondan, {data.get('beruvchi_fio', '___')} (pasport: {data.get('beruvchi_pasport', '___')}), "
        f"bundan buyon «Ijara Beruvchi» deb yuritiladi,"
    )
    doc.add_paragraph(
        f"Ikkinchi tomondan, {data.get('oluvchi_fio', '___')} (pasport: {data.get('oluvchi_pasport', '___')}), "
        f"bundan buyon «Ijara Oluvchi» deb yuritiladi, quyidagi shartnomani tuzdilar:"
    )

    doc.add_paragraph()

    # Moddalar
    doc.add_heading("1. SHARTNOMA PREDMETI", level=2)
    doc.add_paragraph(
        f"1.1. Ijara Beruvchi Ijara Oluvchiga quyidagi manzilda joylashgan ko'chmas mulkni "
        f"vaqtincha foydalanishga beradi: {data.get('manzil', '___')}."
    )

    doc.add_heading("2. IJARA MUDDATI", level=2)
    doc.add_paragraph(
        f"2.1. Shartnoma muddati: {data.get('muddat', '12')} oy. "
        f"Boshlanish sanasi: {data.get('boshlanish', '___')}."
    )

    doc.add_heading("3. IJARA HAQI", level=2)
    doc.add_paragraph(
        f"3.1. Oylik ijara haqi: {data.get('narx', '___')} so'm."
    )
    doc.add_paragraph(
        "3.2. To'lov har oyning 5-kunigacha amalga oshiriladi."
    )

    doc.add_heading("4. TOMONLARNING MAJBURIYATLARI", level=2)
    doc.add_paragraph("4.1. Ijara Beruvchi mulkni yaxshi holatda topshirishga majbur.")
    doc.add_paragraph("4.2. Ijara Oluvchi mulkni ehtiyotkorlik bilan ishlatishga majbur.")

    doc.add_heading("5. IMZOLAR", level=2)
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "IJARA BERUVCHI:"
    table.cell(0, 1).text = "IJARA OLUVCHI:"
    table.cell(1, 0).text = data.get('beruvchi_fio', '___')
    table.cell(1, 1).text = data.get('oluvchi_fio', '___')
    table.cell(2, 0).text = "Imzo: __________"
    table.cell(2, 1).text = "Imzo: __________"

    file_name = f"{TEMPLATES_DIR}/ijara_{order_id}.docx"
    doc.save(file_name)
    return file_name


def generate_oldi_sotdi(data: dict, order_id: int) -> str:
    """Oldi-sotdi shartnomasi generatsiya qilish"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OLDI-SOTDI SHARTNOMASI")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"{data.get('shahar', 'Toshkent')} shahri").bold = True
    p.add_run(f"\t\t\t\t{data.get('sana', datetime.now().strftime('%d.%m.%Y'))} yil")

    doc.add_paragraph()
    doc.add_paragraph(
        f"Sotuvchi: {data.get('sotuvchi_fio', '___')} (pasport: {data.get('sotuvchi_pasport', '___')})"
    )
    doc.add_paragraph(
        f"Xaridor: {data.get('xaridor_fio', '___')} (pasport: {data.get('xaridor_pasport', '___')})"
    )

    doc.add_heading("1. SHARTNOMA PREDMETI", level=2)
    doc.add_paragraph(
        f"1.1. Sotuvchi Xaridorga quyidagi mulkni sotadi: {data.get('tovar', '___')}."
    )

    doc.add_heading("2. NARXI VA TO'LOV", level=2)
    doc.add_paragraph(
        f"2.1. Mulkning narxi: {data.get('narx', '___')} so'm (so'mda)."
    )
    doc.add_paragraph("2.2. To'lov shartnoma imzolanish kuni amalga oshiriladi.")

    doc.add_heading("3. TOMONLAR IMZOLARI", level=2)
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "SOTUVCHI:"
    table.cell(0, 1).text = "XARIDOR:"
    table.cell(1, 0).text = data.get('sotuvchi_fio', '___')
    table.cell(1, 1).text = data.get('xaridor_fio', '___')
    table.cell(2, 0).text = "Imzo: __________"
    table.cell(2, 1).text = "Imzo: __________"

    file_name = f"{TEMPLATES_DIR}/oldi_sotdi_{order_id}.docx"
    doc.save(file_name)
    return file_name


def generate_tilxat(data: dict, order_id: int) -> str:
    """Tilxat generatsiya qilish"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TILXAT")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"{data.get('shahar', 'Toshkent')} shahri, {data.get('sana', datetime.now().strftime('%d.%m.%Y'))} yil")

    doc.add_paragraph()
    doc.add_paragraph(
        f"Men, {data.get('fio', '___')}, pasport seriyasi {data.get('pasport', '___')}, "
        f"{data.get('manzil', '___')} manzilida yashayman."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"Ushbu tilxat bilan tasdiqlayman: {data.get('mazmun', '___')}."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"Miqdori: {data.get('miqdor', '___')} so'm."
    )
    doc.add_paragraph()
    doc.add_paragraph("Imzo: __________")
    doc.add_paragraph(f"Sana: {data.get('sana', datetime.now().strftime('%d.%m.%Y'))}")

    file_name = f"{TEMPLATES_DIR}/tilxat_{order_id}.docx"
    doc.save(file_name)
    return file_name
